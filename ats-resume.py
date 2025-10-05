# -*- coding: utf-8 -*-
"""
Created on Sun Oct  5 10:40:33 2025

@author: deepj
"""

import random
from docx import Document
from docx2pdf import convert
from faker import Faker
import pandas as pd
import os

# Initialize Faker
fake = Faker()

# Directory to save resumes
output_dir = "ATS_Resumes"
os.makedirs(output_dir, exist_ok=True)

# Load job roles and keywords from CSV
# CSV format: Role,Keyword1,Keyword2,Keyword3,Keyword4,Keyword5
csv_file = "job_roles_keywords.csv"  # Make sure this CSV exists
job_data = pd.read_csv(csv_file)

# Function to generate random experience
def generate_experience(role, keywords):
    num_jobs = random.randint(1, 3)
    experiences = []
    for _ in range(num_jobs):
        company = fake.company()
        duration = f"{random.randint(1,5)} years"
        description = f"Worked as {role} at {company} in {random.choice(['IT','Finance','Healthcare','E-commerce','Education','Manufacturing'])} industry. Utilized {', '.join(random.sample(keywords, min(3,len(keywords))))} to achieve key goals and improve efficiency."
        experiences.append({"title": role, "company": company, "duration": duration, "description": description})
    return experiences

# Function to create a resume
def create_resume(file_name, role, keywords):
    doc = Document()
    
    # Contact Info
    doc.add_heading(fake.name(), 0)
    doc.add_paragraph(f"Email: {fake.email()}")
    doc.add_paragraph(f"Phone: {fake.phone_number()}")
    doc.add_paragraph(f"Address: {fake.address()}")
    
    # Professional Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(f"{fake.sentence(nb_words=20)} Experienced {role} with expertise in {', '.join(random.sample(keywords, min(3,len(keywords))))}.")
    
    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(random.sample(keywords, min(5,len(keywords)))))
    
    # Experience
    doc.add_heading("Experience", level=1)
    for exp in generate_experience(role, keywords):
        doc.add_heading(f"{exp['title']} at {exp['company']} ({exp['duration']})", level=2)
        doc.add_paragraph(exp['description'])
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"{fake.university()} - {random.choice(['B.Sc', 'B.Tech', 'M.Sc', 'MBA'])} in {random.choice(['Computer Science', 'Business', 'Engineering', 'Data Science'])}")
    
    # Save Word doc
    doc_path = os.path.join(output_dir, file_name)
    doc.save(doc_path)
    
    # Convert to PDF
    convert(doc_path)

# Generate 100 resumes
for i in range(1, 101):
    row = job_data.sample(n=1).iloc[0]  # Randomly pick a role from CSV
    role = row['Role']
    keywords = [str(row[k]).strip() for k in row.index if k != 'Role' and pd.notna(row[k])]
    create_resume(f"Resume_{i}_{role.replace(' ','_')}.docx", role, keywords)

print("100 ATS-optimized Word and PDF resumes generated successfully in the 'ATS_Resumes' folder.")
